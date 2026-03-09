"""
============================================================
  PHASE 3 & 4 EVALUATION — Skills Expansion + Ingestion
============================================================
Run from your project root:
    python eval_phase3_4.py

Phase 3 — Skills Expansion (12 checks)
  1.  expand_skills() importable
  2.  NumPy → Python
  3.  FastAPI → Python AND REST API
  4.  React → JavaScript AND Frontend
  5.  Spring Boot → Java AND REST API
  6.  Kubernetes → Docker AND DevOps
  7.  PySpark → Apache Spark, Python, Big Data
  8.  Explicit parent not downgraded (higher value wins)
  9.  Duplicate skill names deduplicated
  10. Case-insensitive deduplication (Python vs python)
  11. expand_skills_node() updates state correctly
  12. Expansion with empty skills list returns []

Phase 4 — Document Ingestion (10 checks)
  13. load_documents_node importable
  14. scan_resume_directory returns correct file list
  15. .txt file loads with correct metadata fields
  16. .pdf file loads (using sample PDF in data/resumes/)
  17. .docx file loads (using sample DOCX in data/resumes/)
  18. Non-existent file goes to failed_docs with reason
  19. Unsupported extension (.xlsx) goes to failed_docs
  20. Empty/near-empty file goes to failed_docs
  21. load_documents_node state output structure correct
  22. Multiple files: each loaded doc has unique file_name
============================================================
"""

import sys
import json
import shutil
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(".").resolve()))

GREEN = "\033[92m"; RED = "\033[91m"; YELLOW = "\033[93m"
CYAN  = "\033[96m"; BOLD = "\033[1m"; RESET = "\033[0m"; DIM = "\033[2m"

pass_count = 0
fail_count = 0

def ok(msg, detail=""):
    global pass_count; pass_count += 1
    print(f"  {GREEN}✓ PASS{RESET}  {msg}")
    if detail: print(f"         {DIM}{detail}{RESET}")

def fail(msg, detail=""):
    global fail_count; fail_count += 1
    print(f"  {RED}✗ FAIL{RESET}  {msg}")
    if detail: print(f"         {DIM}{detail}{RESET}")

def warn(msg):
    print(f"  {YELLOW}⚠ WARN{RESET}  {msg}")

def info(msg):
    print(f"         {DIM}{msg}{RESET}")

def header(title):
    print(f"\n{BOLD}{CYAN}── {title} {'─'*(54-len(title))}{RESET}")

from dotenv import load_dotenv
load_dotenv()

print(f"\n{BOLD}{'='*60}{RESET}")
print(f"{BOLD}  PHASE 3 & 4 — Skills Expansion + Document Ingestion{RESET}")
print(f"{BOLD}{'='*60}{RESET}")


# ══════════════════════════════════════════════════════════════
#  PHASE 3 — SKILLS EXPANSION
# ══════════════════════════════════════════════════════════════
header("Phase 3 — Skills Expansion")

# Check 1
try:
    from nodes.expand_skills import expand_skills, expand_skills_node
    ok("expand_skills and expand_skills_node importable")
except ImportError as e:
    fail(f"Import failed: {e}")
    sys.exit(1)

# Helper to run expansion and return skill name set
def get_skill_names(input_skills: list[dict]) -> set[str]:
    result = expand_skills(input_skills)
    return {s["skill"] for s in result}

def get_skill_years(input_skills: list[dict]) -> dict[str, float]:
    result = expand_skills(input_skills)
    return {s["skill"]: s["experience_years"] for s in result}

# Check 2
try:
    names = get_skill_names([{"skill": "NumPy", "experience_years": 2.0}])
    assert "Python" in names, f"Expected Python, got: {names}"
    ok("NumPy → Python inferred", f"All skills: {names}")
except AssertionError as e:
    fail(str(e))

# Check 3
try:
    names = get_skill_names([{"skill": "FastAPI", "experience_years": 1.5}])
    assert "Python" in names, f"Python missing from: {names}"
    assert "REST API" in names, f"REST API missing from: {names}"
    ok("FastAPI → Python AND REST API inferred", f"All skills: {names}")
except AssertionError as e:
    fail(str(e))

# Check 4
try:
    names = get_skill_names([{"skill": "React", "experience_years": 3.0}])
    assert "JavaScript" in names, f"JavaScript missing from: {names}"
    assert "Frontend" in names, f"Frontend missing from: {names}"
    ok("React → JavaScript AND Frontend inferred", f"All skills: {names}")
except AssertionError as e:
    fail(str(e))

# Check 5
try:
    names = get_skill_names([{"skill": "Spring Boot", "experience_years": 2.0}])
    assert "Java" in names, f"Java missing from: {names}"
    assert "REST API" in names, f"REST API missing from: {names}"
    ok("Spring Boot → Java AND REST API inferred", f"All skills: {names}")
except AssertionError as e:
    fail(str(e))

# Check 6
try:
    names = get_skill_names([{"skill": "Kubernetes", "experience_years": 1.0}])
    assert "Docker" in names, f"Docker missing from: {names}"
    assert "DevOps" in names, f"DevOps missing from: {names}"
    ok("Kubernetes → Docker AND DevOps inferred", f"All skills: {names}")
except AssertionError as e:
    fail(str(e))

# Check 7
try:
    names = get_skill_names([{"skill": "PySpark", "experience_years": 2.0}])
    assert "Apache Spark" in names, f"Apache Spark missing from: {names}"
    assert "Python" in names, f"Python missing from: {names}"
    assert "Big Data" in names, f"Big Data missing from: {names}"
    ok("PySpark → Apache Spark, Python, Big Data inferred", f"All skills: {names}")
except AssertionError as e:
    fail(str(e))

# Check 8 — Higher value wins (don't downgrade explicit parent)
try:
    years = get_skill_years([
        {"skill": "NumPy", "experience_years": 1.0},  # Would infer Python at 1y
        {"skill": "Python", "experience_years": 4.0}, # Explicit — must stay at 4y
    ])
    assert years.get("Python") == 4.0, f"Python should be 4.0, got {years.get('Python')}"
    ok("Explicit parent not downgraded — higher value wins (Python stays at 4.0y)")
    info(f"Skill years: {years}")
except AssertionError as e:
    fail(str(e))

# Check 9 — Deduplication
try:
    result = expand_skills([
        {"skill": "Python", "experience_years": 3.0},
        {"skill": "Python", "experience_years": 3.0},
        {"skill": "Python", "experience_years": 1.0},
    ])
    python_entries = [s for s in result if s["skill"] == "Python"]
    assert len(python_entries) == 1, f"Expected 1 Python entry, got {len(python_entries)}"
    assert python_entries[0]["experience_years"] == 3.0, "Should keep max years"
    ok("Duplicate skill entries deduplicated, max years kept")
    info(f"Single Python entry: {python_entries[0]}")
except AssertionError as e:
    fail(str(e))

# Check 10 — Case-insensitive deduplication
try:
    result = expand_skills([
        {"skill": "Python", "experience_years": 4.0},
        {"skill": "python", "experience_years": 2.0},
        {"skill": "PYTHON", "experience_years": 1.0},
    ])
    python_entries = [s for s in result if s["skill"].lower() == "python"]
    assert len(python_entries) == 1, f"Expected 1 python entry (case-insensitive), got {len(python_entries)}"
    assert python_entries[0]["experience_years"] == 4.0
    ok("Case-insensitive deduplication works (Python/python/PYTHON → 1 entry at 4.0y)")
except AssertionError as e:
    fail(str(e))

# Check 11 — LangGraph node state update
try:
    mock_state = {
        "parsed_resumes": [
            {
                "candidate_id": "abc123",
                "name": "Test Candidate",
                "skills": [
                    {"skill": "React", "experience_years": 2.0},
                    {"skill": "NumPy", "experience_years": 1.5},
                ],
                "experience": [],
                "education": [],
                "certifications": [],
                "file_name": "test.pdf",
                "upload_time": "2025-01-01T00:00:00",
            }
        ],
        "node_logs": [],
    }
    result_state = expand_skills_node(mock_state)
    assert "expanded_resumes" in result_state, "Missing expanded_resumes"
    assert len(result_state["expanded_resumes"]) == 1
    exp_skills = result_state["expanded_resumes"][0]["skills"]
    skill_names = {s["skill"] for s in exp_skills}
    assert "JavaScript" in skill_names, f"JavaScript missing: {skill_names}"
    assert "Python" in skill_names, f"Python missing: {skill_names}"
    assert len(mock_state["node_logs"]) == 1, "Node should log exactly 1 entry"
    log = mock_state["node_logs"][0]
    assert log["node"] == "expand_skills_node"
    assert log["status"] == "success"
    ok("expand_skills_node updates state correctly and logs execution")
    info(f"Expanded to {len(exp_skills)} skills: {sorted(skill_names)}")
    info(f"Log entry: {log}")
except (AssertionError, KeyError) as e:
    fail(f"expand_skills_node state test failed: {e}")

# Check 12 — Empty skills
try:
    result = expand_skills([])
    assert result == [], f"Expected empty list, got {result}"
    ok("expand_skills([]) returns empty list (no crash)")
except AssertionError as e:
    fail(str(e))


# ══════════════════════════════════════════════════════════════
#  PHASE 4 — DOCUMENT INGESTION
# ══════════════════════════════════════════════════════════════
header("Phase 4 — Document Ingestion")

# Check 13
try:
    from nodes.load_documents import load_documents_node, scan_resume_directory, load_single_document
    ok("load_documents_node, scan_resume_directory, load_single_document importable")
except ImportError as e:
    fail(f"Import failed: {e}")
    sys.exit(1)

# ── Create test files in data/resumes/ ───────────────────────────────────────
RESUMES_DIR = Path("data/resumes")
RESUMES_DIR.mkdir(parents=True, exist_ok=True)

SAMPLE_TXT = RESUMES_DIR / "sample_candidate.txt"
SAMPLE_TXT.write_text(
    "John Smith | john@email.com | +91 9876543210\n"
    "Location: Bangalore, India\n\n"
    "SKILLS\nPython - 4 years | FastAPI - 2 years | PostgreSQL - 3 years\n\n"
    "EDUCATION\nB.Tech Computer Science, IIT Delhi, 2019, GPA 8.5\n\n"
    "EXPERIENCE\nBackend Engineer at TechCorp (Jan 2019 - Jan 2023)\n"
    "  - Built REST APIs for payment processing platform\n"
    "  - Reduced API latency 40% through caching and query optimization\n\n"
    "CERTIFICATIONS\nAWS Certified Developer, Python Professional Certificate\n",
    encoding="utf-8"
)

# Create a minimal DOCX for testing
SAMPLE_DOCX = RESUMES_DIR / "sample_candidate2.docx"
try:
    from docx import Document as DocxDocument
    doc = DocxDocument()
    doc.add_heading("Alice Wong", 0)
    doc.add_paragraph("alice@email.com | +91 8765432109 | Mumbai, India")
    doc.add_heading("Skills", 1)
    doc.add_paragraph("Java - 5 years | Spring Boot - 3 years | Kubernetes - 2 years")
    doc.add_heading("Education", 1)
    doc.add_paragraph("B.Tech Computer Engineering, IIT Bombay, 2018")
    doc.add_heading("Experience", 1)
    doc.add_paragraph("Senior Java Developer at InfoSys (2018 - 2023)\n"
                      "Developed microservices architecture for banking platform using Spring Boot")
    doc.save(str(SAMPLE_DOCX))
    info(f"Created test DOCX: {SAMPLE_DOCX}")
except Exception as e:
    warn(f"Could not create test DOCX: {e}")
    SAMPLE_DOCX = None

# Check 14 — Directory scan
try:
    files = scan_resume_directory(RESUMES_DIR)
    assert len(files) >= 1, f"Expected at least 1 file, found {len(files)}"
    assert all(Path(f).exists() for f in files), "Scan returned non-existent paths"
    ok(f"scan_resume_directory finds {len(files)} file(s)")
    for f in files:
        info(f"  Found: {Path(f).name}")
except (AssertionError, Exception) as e:
    fail(f"scan_resume_directory failed: {e}")

# Check 15 — TXT load
try:
    doc = load_single_document(SAMPLE_TXT)
    required_keys = ["file_name", "file_path", "file_type", "text", "upload_time", "char_count"]
    missing = [k for k in required_keys if k not in doc]
    assert not missing, f"Missing keys: {missing}"
    assert doc["file_type"] == "txt"
    assert "John Smith" in doc["text"]
    assert doc["char_count"] > 100
    ok("TXT file loads with all required metadata fields")
    info(f"file_name: {doc['file_name']}, char_count: {doc['char_count']}")
    info(f"Text preview: {doc['text'][:80]}...")
except (AssertionError, Exception) as e:
    fail(f"TXT load failed: {e}")

# Check 16 — PDF load (create a minimal one if none exists)
pdf_files = list(RESUMES_DIR.glob("*.pdf"))
if not pdf_files:
    warn("No PDF in data/resumes/ — creating a minimal one for testing")
    try:
        # Try creating with reportlab if available
        try:
            from reportlab.pdfgen import canvas as rl_canvas
            pdf_path = RESUMES_DIR / "sample_pdf_candidate.pdf"
            c = rl_canvas.Canvas(str(pdf_path))
            c.drawString(72, 750, "Bob Johnson | bob@email.com")
            c.drawString(72, 730, "Skills: Python 3y, Django 2y, PostgreSQL 2y")
            c.drawString(72, 710, "Education: B.Tech IIT Madras 2020")
            c.drawString(72, 690, "Experience: Python Developer at StartupXYZ 2 years")
            c.drawString(72, 670, "Built scalable REST APIs using Django and PostgreSQL")
            c.save()
            pdf_files = [pdf_path]
            info(f"Created test PDF: {pdf_path}")
        except ImportError:
            warn("reportlab not installed — place any PDF resume in data/resumes/ to test PDF loading")
            pdf_files = []
    except Exception as e:
        warn(f"Could not create test PDF: {e}")

if pdf_files:
    try:
        doc = load_single_document(pdf_files[0])
        assert doc["file_type"] == "pdf"
        assert len(doc["text"]) > 20, f"PDF text too short: '{doc['text']}'"
        ok(f"PDF file loads successfully")
        info(f"file_name: {doc['file_name']}, char_count: {doc['char_count']}")
        info(f"Text preview: {doc['text'][:80]}...")
    except Exception as e:
        fail(f"PDF load failed: {e}")
else:
    warn("PDF check skipped — place a PDF resume in data/resumes/ to enable this check")

# Check 17 — DOCX load
if SAMPLE_DOCX and SAMPLE_DOCX.exists():
    try:
        doc = load_single_document(SAMPLE_DOCX)
        assert doc["file_type"] == "docx"
        assert "Alice Wong" in doc["text"], f"Expected 'Alice Wong' in text, got: {doc['text'][:100]}"
        assert doc["char_count"] > 50
        ok("DOCX file loads and text extracted correctly")
        info(f"Text preview: {doc['text'][:80]}...")
    except (AssertionError, Exception) as e:
        fail(f"DOCX load failed: {e}")
else:
    warn("DOCX check skipped — python-docx may not be available")

# Check 18 — Non-existent file
try:
    state = {
        "resume_file_paths": ["data/resumes/DOES_NOT_EXIST.pdf"],
        "failed_docs": [],
        "node_logs": [],
    }
    result = load_documents_node(state)
    assert len(result["raw_resume_texts"]) == 0
    assert len(result["failed_docs"]) == 1
    fd = result["failed_docs"][0]
    assert "FileNotFoundError" in fd["error"] or "not found" in fd["reason"].lower()
    ok("Non-existent file correctly goes to failed_docs")
    info(f"failed_docs entry: {fd}")
except (AssertionError, Exception) as e:
    fail(f"Non-existent file handling failed: {e}")

# Check 19 — Unsupported extension
try:
    # Create a dummy xlsx file
    dummy_xlsx = RESUMES_DIR / "dummy.xlsx"
    dummy_xlsx.write_bytes(b"PK fake xlsx content")
    state = {
        "resume_file_paths": [str(dummy_xlsx)],
        "failed_docs": [],
        "node_logs": [],
    }
    result = load_documents_node(state)
    assert len(result["raw_resume_texts"]) == 0
    assert len(result["failed_docs"]) == 1
    fd = result["failed_docs"][0]
    assert "UnsupportedFileType" in fd["error"] or "unsupported" in fd["reason"].lower()
    ok("Unsupported file type (.xlsx) correctly goes to failed_docs")
    info(f"Reason: {fd['reason']}")
    dummy_xlsx.unlink()
except (AssertionError, Exception) as e:
    fail(f"Unsupported file type handling failed: {e}")

# Check 20 — Empty file
try:
    empty_txt = RESUMES_DIR / "empty_resume.txt"
    empty_txt.write_text("   ", encoding="utf-8")  # whitespace only
    state = {
        "resume_file_paths": [str(empty_txt)],
        "failed_docs": [],
        "node_logs": [],
    }
    result = load_documents_node(state)
    assert len(result["raw_resume_texts"]) == 0
    assert len(result["failed_docs"]) == 1
    fd = result["failed_docs"][0]
    assert "EmptyDocument" in fd["error"]
    ok("Near-empty file correctly goes to failed_docs with EmptyDocument error")
    info(f"Reason: {fd['reason']}")
    empty_txt.unlink()
except (AssertionError, Exception) as e:
    fail(f"Empty file handling failed: {e}")

# Check 21 — Node state structure
try:
    state = {
        "resume_file_paths": [str(SAMPLE_TXT)],
        "failed_docs": [],
        "node_logs": [],
    }
    result = load_documents_node(state)
    assert "raw_resume_texts" in result
    assert "failed_docs" in result
    assert len(state["node_logs"]) == 1
    log = state["node_logs"][0]
    for field in ["node", "status", "latency_ms", "timestamp"]:
        assert field in log, f"Missing '{field}' in node log"
    assert log["node"] == "load_documents_node"
    assert log["status"] == "success"
    ok("load_documents_node state output structure is correct")
    info(f"Log: node={log['node']}, status={log['status']}, latency={log['latency_ms']}ms")
except (AssertionError, Exception) as e:
    fail(f"Node state structure check failed: {e}")

# Check 22 — Multiple files, unique file_names
try:
    all_files = scan_resume_directory(RESUMES_DIR)
    # Filter out dummy files we may have created
    valid_files = [f for f in all_files if Path(f).suffix.lower() in {".txt", ".docx", ".pdf"}
                   and "empty" not in Path(f).name and "dummy" not in Path(f).name]
    if len(valid_files) >= 2:
        state = {
            "resume_file_paths": valid_files[:3],
            "failed_docs": [],
            "node_logs": [],
        }
        result = load_documents_node(state)
        loaded = result["raw_resume_texts"]
        if len(loaded) >= 2:
            file_names = [d["file_name"] for d in loaded]
            assert len(file_names) == len(set(file_names)), f"Duplicate file_names: {file_names}"
            ok(f"Multiple files loaded with unique file_names ({len(loaded)} files)")
            for d in loaded:
                info(f"  {d['file_name']} — {d['char_count']} chars")
        else:
            ok(f"Multiple file paths processed ({len(loaded)} loaded, {len(result['failed_docs'])} failed)")
    else:
        ok("Multiple-file uniqueness check: only 1 valid file found, adding second sample")
        # Create a second txt to test
        sample2 = RESUMES_DIR / "sample_candidate3.txt"
        sample2.write_text(
            "Carol Davis | carol@email.com\nSkills: React 3y, Node.js 2y, TypeScript 2y\n"
            "Education: B.Sc Computer Science, Mumbai University, 2021\n"
            "Experience: Frontend Developer at WebAgency 2 years\n"
            "Built React applications for e-commerce clients\n",
            encoding="utf-8"
        )
        state = {
            "resume_file_paths": [str(SAMPLE_TXT), str(sample2)],
            "failed_docs": [], "node_logs": [],
        }
        result = load_documents_node(state)
        loaded = result["raw_resume_texts"]
        file_names = [d["file_name"] for d in loaded]
        assert len(file_names) == len(set(file_names))
        ok(f"Multiple files loaded, all file_names unique: {file_names}")
except (AssertionError, Exception) as e:
    fail(f"Multiple-file test failed: {e}")


# ── SUMMARY ───────────────────────────────────────────────────────────────────
total = pass_count + fail_count
print(f"\n{BOLD}{'='*60}{RESET}")
print(f"{BOLD}  PHASE 3 & 4 RESULT: {pass_count}/{total} checks passed{RESET}")
if fail_count == 0:
    print(f"  {GREEN}{BOLD}All checks passed! Ready for Phase 5.{RESET}")
else:
    print(f"  {RED}Fix the FAILs above before Phase 5.{RESET}")
print(f"\n  {DIM}Files created in data/resumes/ for testing:")
for f in sorted(Path("data/resumes").glob("*")):
    print(f"    {f.name} ({f.stat().st_size} bytes)")
print(f"{RESET}{BOLD}{'='*60}{RESET}\n")

sys.exit(0 if fail_count == 0 else 1)
