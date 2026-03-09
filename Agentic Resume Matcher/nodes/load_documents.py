"""
nodes/load_documents.py
-----------------------
Document Ingestion Node — LangGraph node.

Loads resume files from disk and converts them to plain text with metadata.
Supports: PDF (via pdfplumber), DOCX (via python-docx), TXT (plain read).

Each loaded document becomes:
{
    "file_name":   "jane_doe.pdf",
    "file_path":   "/abs/path/to/file",
    "file_type":   "pdf",               # "pdf" | "docx" | "txt"
    "text":        "Jane Doe | Python...",
    "upload_time": "2025-01-01T10:30:00",
    "char_count":  1842,
}

Documents with empty text or load errors are skipped and logged to failed_docs.
"""

import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from utils.logger import NodeTimer


# ── Loader functions (one per file type) ─────────────────────────────────────

def _load_pdf(file_path: Path) -> str:
    """Extract text from PDF using pdfplumber (better than pypdf for layout)."""
    import pdfplumber
    pages = []
    with pdfplumber.open(str(file_path)) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text.strip())
    return "\n".join(pages)


def _load_docx(file_path: Path) -> str:
    """Extract text from DOCX, including tables."""
    from docx import Document
    doc = Document(str(file_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # Also extract table cells — resumes often use tables for layout
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text and text not in paragraphs:
                    paragraphs.append(text)
    return "\n".join(paragraphs)


def _load_txt(file_path: Path) -> str:
    """Load plain text file."""
    return file_path.read_text(encoding="utf-8", errors="replace").strip()


_LOADERS = {
    ".pdf":  _load_pdf,
    ".docx": _load_docx,
    ".doc":  _load_docx,   # python-docx handles .doc too in most cases
    ".txt":  _load_txt,
    ".text": _load_txt,
}

SUPPORTED_EXTENSIONS = set(_LOADERS.keys())


# ── Core load function (pure, testable independently) ─────────────────────────

def load_single_document(file_path: str | Path) -> dict[str, Any]:
    """
    Load a single file and return a document dict.

    Returns a dict with keys: file_name, file_path, file_type, text,
                               upload_time, char_count
    Raises ValueError if file type is unsupported.
    Raises any IO/parsing exception from the underlying loader.
    """
    path = Path(file_path)
    suffix = path.suffix.lower()

    if suffix not in _LOADERS:
        raise ValueError(
            f"Unsupported file type '{suffix}'. "
            f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    loader = _LOADERS[suffix]
    text = loader(path)

    return {
        "file_name":   path.name,
        "file_path":   str(path.resolve()),
        "file_type":   suffix.lstrip("."),
        "text":        text,
        "upload_time": datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S"),
        "char_count":  len(text),
    }


# ── LangGraph node ────────────────────────────────────────────────────────────

def load_documents_node(state: dict) -> dict:
    """
    LangGraph node: Load all resume files into text documents.

    Reads:  state["resume_file_paths"] — list of file path strings
    Writes: state["raw_resume_texts"]  — list of document dicts
            state["failed_docs"]       — appends load failures

    Files that fail to load are added to failed_docs with a reason,
    rather than crashing the whole pipeline.
    """
    file_paths: list[str] = state.get("resume_file_paths", [])
    existing_failed: list[dict] = state.get("failed_docs", [])

    with NodeTimer("load_documents_node", state) as timer:
        loaded_docs = []
        failed_docs = list(existing_failed)
        skipped_empty = 0

        for file_path in file_paths:
            path = Path(file_path)

            # Check existence
            if not path.exists():
                failed_docs.append({
                    "file_name": path.name,
                    "error":     "FileNotFoundError",
                    "reason":    f"File not found: {file_path}",
                    "stage":     "load_documents_node",
                })
                continue

            # Check extension
            if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                failed_docs.append({
                    "file_name": path.name,
                    "error":     "UnsupportedFileType",
                    "reason":    f"Unsupported extension: {path.suffix}",
                    "stage":     "load_documents_node",
                })
                continue

            # Load
            try:
                doc = load_single_document(path)

                # Skip empty files
                if len(doc["text"].strip()) < 50:
                    skipped_empty += 1
                    failed_docs.append({
                        "file_name": path.name,
                        "error":     "EmptyDocument",
                        "reason":    f"Extracted text too short ({doc['char_count']} chars). Likely a scanned/image PDF.",
                        "stage":     "load_documents_node",
                    })
                    continue

                loaded_docs.append(doc)

            except Exception as e:
                failed_docs.append({
                    "file_name": path.name,
                    "error":     type(e).__name__,
                    "reason":    str(e),
                    "stage":     "load_documents_node",
                })

        timer.extra = {
            "files_requested": len(file_paths),
            "files_loaded":    len(loaded_docs),
            "files_failed":    len(failed_docs) - len(existing_failed),
            "skipped_empty":   skipped_empty,
        }

    return {
        "raw_resume_texts": loaded_docs,
        "failed_docs":      failed_docs,
    }


# ── Directory scanner helper ──────────────────────────────────────────────────

def scan_resume_directory(directory: str | Path) -> list[str]:
    """
    Utility: scan a directory and return paths of all supported resume files.
    Used by the Streamlit UI and ingestion trigger.
    """
    dir_path = Path(directory)
    if not dir_path.exists():
        return []
    return [
        str(f) for f in sorted(dir_path.iterdir())
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]