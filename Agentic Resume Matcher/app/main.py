# -*- coding: utf-8 -*-
"""
app/main.py  --  Agentic Resume Matcher
Run:   streamlit run app/main.py
Login: admin / hr2025   or   recruiter / match123
"""

import sys
import os
import warnings
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
# Suppress Pydantic serializer warnings globally -- these come from LangChain
# internals when a Pydantic v2 model is passed through LangGraph state.
warnings.filterwarnings("ignore", message=".*PydanticSerializationUnexpectedValue.*")
os.environ["PYTHONWARNINGS"] = "ignore::UserWarning"

import streamlit as st

st.set_page_config(
    page_title="Resume Matcher",
    page_icon=":mag:",
    layout="wide",
    initial_sidebar_state="expanded",
)

CSS = (
    "<style>"
    "@import url('https://fonts.googleapis.com/css2?family=DM+Mono:wght@400;500"
    "&family=Syne:wght@400;700;800&display=swap');"
    ":root{"
    "--bg:#0d0f14;--surface:#13151c;--border:#1e2130;"
    "--accent:#5b6af5;--accent2:#38d9a9;--danger:#f06595;"
    "--text:#e8eaf0;--muted:#6b7280;"
    "}"
    "html,body,[class*='css']{"
    "background:var(--bg)!important;color:var(--text)!important;"
    "font-family:'Syne',sans-serif!important;"
    "}"
    "[data-testid='stSidebar']{"
    "background:var(--surface)!important;border-right:1px solid var(--border)!important;"
    "}"
    ".top-bar{display:flex;align-items:center;gap:12px;"
    "padding:0 0 20px;border-bottom:1px solid var(--border);margin-bottom:24px;}"
    ".top-bar h1{font-family:'Syne',sans-serif;font-size:1.35rem;font-weight:800;"
    "letter-spacing:-0.02em;color:var(--text);margin:0;}"
    ".badge{font-family:'DM Mono',monospace;font-size:0.68rem;color:var(--accent2);"
    "background:rgba(56,217,169,0.08);border:1px solid rgba(56,217,169,0.18);"
    "border-radius:4px;padding:2px 8px;margin-left:auto;}"
    ".shdr{font-family:'DM Mono',monospace;font-size:0.68rem;color:var(--muted);"
    "text-transform:uppercase;letter-spacing:0.12em;"
    "border-bottom:1px solid var(--border);padding-bottom:7px;margin:18px 0 12px;}"
    ".ccard{background:var(--surface);border:1px solid var(--border);"
    "border-radius:10px;padding:16px 18px;margin-bottom:10px;"
    "position:relative;overflow:hidden;}"
    ".ccard::before{content:'';position:absolute;left:0;top:0;bottom:0;"
    "width:3px;background:var(--accent);border-radius:10px 0 0 10px;}"
    ".ccard.r1::before{background:var(--accent2);}"
    ".cname{font-family:'Syne',sans-serif;font-size:0.95rem;font-weight:700;"
    "color:var(--text);margin-bottom:3px;}"
    ".csub{font-family:'DM Mono',monospace;font-size:0.62rem;"
    "color:var(--muted);text-transform:uppercase;letter-spacing:0.1em;}"
    ".sbwrap{margin:7px 0 3px;}"
    ".sblbl{font-family:'DM Mono',monospace;font-size:0.62rem;color:var(--muted);"
    "text-transform:uppercase;display:flex;justify-content:space-between;"
    "margin-bottom:3px;}"
    ".sb{height:4px;background:var(--border);border-radius:2px;overflow:hidden;}"
    ".sf{height:100%;border-radius:2px;"
    "background:linear-gradient(90deg,var(--accent),var(--accent2));}"
    ".sf2{height:100%;border-radius:2px;"
    "background:linear-gradient(90deg,var(--accent2),var(--accent));}"
    ".chip{display:inline-block;font-family:'DM Mono',monospace;font-size:0.62rem;"
    "background:rgba(91,106,245,0.1);color:var(--accent);"
    "border:1px solid rgba(91,106,245,0.22);border-radius:4px;"
    "padding:2px 7px;margin:2px 2px 2px 0;}"
    ".chip.m{background:rgba(56,217,169,0.09);color:var(--accent2);"
    "border-color:rgba(56,217,169,0.22);}"
    ".mrow{display:flex;gap:10px;margin-top:10px;}"
    ".mbox{flex:1;background:rgba(255,255,255,0.03);border:1px solid var(--border);"
    "border-radius:6px;padding:7px 10px;text-align:center;}"
    ".mval{font-family:'DM Mono',monospace;font-size:0.95rem;font-weight:500;"
    "color:var(--accent2);}"
    ".mlbl{font-family:'DM Mono',monospace;font-size:0.58rem;color:var(--muted);"
    "text-transform:uppercase;letter-spacing:0.08em;margin-top:2px;}"
    ".jpanel{background:var(--surface);border:1px solid var(--border);"
    "border-radius:10px;padding:16px 18px;margin-bottom:16px;}"
    ".jtitle{font-size:0.95rem;font-weight:700;margin-bottom:3px;}"
    ".jmeta{font-family:'DM Mono',monospace;font-size:0.68rem;color:var(--muted);}"
    ".scard{background:var(--surface);border:1px solid var(--border);"
    "border-radius:10px;padding:18px;text-align:center;}"
    ".sval{font-family:'DM Mono',monospace;font-size:1.8rem;font-weight:500;"
    "color:var(--accent2);line-height:1;}"
    ".slbl{font-family:'DM Mono',monospace;font-size:0.62rem;color:var(--muted);"
    "text-transform:uppercase;letter-spacing:0.1em;margin-top:5px;}"
    ".lrow{font-family:'DM Mono',monospace;font-size:0.68rem;color:var(--muted);"
    "padding:3px 0;border-bottom:1px solid var(--border);display:flex;gap:10px;}"
    ".lnode{color:var(--accent);min-width:190px;}"
    ".lms{color:var(--accent2);min-width:65px;}"
    ".lerr{color:var(--danger)!important;}"
    ".stButton>button{font-family:'DM Mono',monospace!important;"
    "font-size:0.73rem!important;background:var(--accent)!important;"
    "color:white!important;border:none!important;border-radius:6px!important;}"
    ".stButton>button:hover{background:#4a59e0!important;}"
    ".stTextArea textarea,.stTextInput input{"
    "background:var(--surface)!important;border:1px solid var(--border)!important;"
    "color:var(--text)!important;font-family:'DM Mono',monospace!important;"
    "font-size:0.78rem!important;border-radius:6px!important;}"
    ".stProgress>div>div{background:var(--accent2)!important;}"
    "#MainMenu,footer,header{visibility:hidden;}"
    "::-webkit-scrollbar{width:4px;}"
    "::-webkit-scrollbar-thumb{background:var(--border);border-radius:2px;}"
    "</style>"
)
st.markdown(CSS, unsafe_allow_html=True)


# ── Auth ───────────────────────────────────────────────────────────────────────
USERS = {"admin": "hr2025", "recruiter": "match123"}


def check_auth():
    if not st.session_state.get("authenticated"):
        col = st.columns([1, 1.2, 1])[1]
        with col:
            st.markdown("<br><br>", unsafe_allow_html=True)
            st.markdown(
                "<div style='text-align:center;margin-bottom:24px;'>"
                "<h2 style='font-family:Syne,sans-serif;font-weight:800;"
                "letter-spacing:-0.02em;'>Resume Matcher</h2>"
                "<p style='font-family:DM Mono,monospace;font-size:0.68rem;"
                "color:#6b7280;'>AGENTIC RECRUITMENT PLATFORM</p></div>",
                unsafe_allow_html=True,
            )
            u = st.text_input("Username", placeholder="username")
            p = st.text_input("Password", type="password", placeholder="password")
            if st.button("Sign In", use_container_width=True):
                if USERS.get(u) == p:
                    st.session_state["authenticated"] = True
                    st.session_state["username"] = u
                    st.rerun()
                else:
                    st.error("Invalid credentials")
        st.stop()


check_auth()


# ── Sidebar ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown(
        "<div style='padding:6px 0 18px;text-align:center;'>"
        "<div style='font-family:Syne,sans-serif;font-weight:800;font-size:0.95rem;'>"
        "Resume Matcher</div>"
        "<div style='font-family:DM Mono,monospace;font-size:0.58rem;color:#6b7280;'>"
        "AGENTIC AI PLATFORM</div></div>",
        unsafe_allow_html=True,
    )

    page = st.radio(
        "nav",
        ["Match Candidates", "Ingest Resumes", "Analytics"],
        label_visibility="collapsed",
    )
    st.markdown("<br>", unsafe_allow_html=True)

    try:
        from config.settings import get_qdrant_client, QDRANT_RESUME_COLLECTION
        _cnt = get_qdrant_client().get_collection(QDRANT_RESUME_COLLECTION).points_count
        st.markdown(
            "<div style='font-family:DM Mono,monospace;font-size:0.63rem;color:#6b7280;"
            "background:#13151c;border:1px solid #1e2130;border-radius:6px;padding:9px 11px;'>"
            "<div style='color:#38d9a9;font-size:0.88rem;font-weight:500;'>"
            + str(_cnt) + "</div>"
            "<div style='margin-top:2px;'>resumes indexed</div></div>",
            unsafe_allow_html=True,
        )
    except Exception:
        pass

    st.markdown("<br><br>", unsafe_allow_html=True)
    if st.button("Sign Out", use_container_width=True):
        st.session_state["authenticated"] = False
        st.rerun()


# ── HTML helpers ──────────────────────────────────────────────────────────────
def sbar(label, value, cls="sf"):
    w   = str(round(value * 100, 1))
    pct = str(round(value * 100)) + "%"
    return (
        "<div class='sbwrap'>"
        "<div class='sblbl'><span>" + label + "</span><span>" + pct + "</span></div>"
        "<div class='sb'><div class='" + cls + "' style='width:" + w + "%;'></div></div>"
        "</div>"
    )


def mbox(val, lbl):
    return (
        "<div class='mbox'>"
        "<div class='mval'>" + str(val) + "</div>"
        "<div class='mlbl'>" + lbl + "</div>"
        "</div>"
    )


def lrow(node, ms, status):
    cls = " lerr" if status == "error" else ""
    return (
        "<div class='lrow" + cls + "'>"
        "<span class='lnode'>" + str(node) + "</span>"
        "<span class='lms'>" + str(round(ms)) + "ms</span>"
        "<span>" + str(status) + "</span>"
        "</div>"
    )


# ==============================================================================
#  PAGE: MATCH CANDIDATES
# ==============================================================================
if page == "Match Candidates":
    st.markdown(
        "<div class='top-bar'><h1>Match Candidates</h1>"
        "<span class='badge'>QUERY PIPELINE</span></div>",
        unsafe_allow_html=True,
    )

    left, right = st.columns([1, 1.6], gap="large")

    with left:
        st.markdown("<div class='shdr'>Job Description</div>", unsafe_allow_html=True)

        jd_src = st.radio("src", ["Paste Text", "Upload File", "From Folder"],
                          horizontal=True, label_visibility="collapsed")
        jd_text  = ""
        jd_fname = None

        if jd_src == "Paste Text":
            jd_text = st.text_area("jd", height=230,
                                   placeholder="Paste job description here...",
                                   label_visibility="collapsed")

        elif jd_src == "Upload File":
            up = st.file_uploader("uf", type=["txt", "pdf", "docx"],
                                  label_visibility="collapsed")
            if up:
                raw = up.read()
                if up.type == "application/pdf":
                    import pdfplumber, io
                    with pdfplumber.open(io.BytesIO(raw)) as pdf:
                        jd_text = "\n".join(p.extract_text() or "" for p in pdf.pages)
                elif up.name.lower().endswith(".docx"):
                    from docx import Document
                    import io
                    doc = Document(io.BytesIO(raw))
                    jd_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                else:
                    jd_text = raw.decode("utf-8", errors="replace")
                jd_fname = up.name
                st.text_area("prev", jd_text[:400] + "...", height=90,
                             disabled=True, label_visibility="collapsed")

        else:
            from config.settings import JD_DIR
            jd_files = (
                sorted(JD_DIR.glob("*.txt"))
                + sorted(JD_DIR.glob("*.pdf"))
                + sorted(JD_DIR.glob("*.docx"))
            )
            if jd_files:
                sel      = st.selectbox("jsel", [f.name for f in jd_files],
                                        label_visibility="collapsed")
                jd_path  = JD_DIR / sel
                jd_fname = sel
                if jd_path.suffix == ".pdf":
                    import pdfplumber
                    with pdfplumber.open(str(jd_path)) as pdf:
                        jd_text = "\n".join(pg.extract_text() or "" for pg in pdf.pages)
                elif jd_path.suffix == ".docx":
                    from docx import Document
                    doc = Document(str(jd_path))
                    jd_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
                else:
                    jd_text = jd_path.read_text(encoding="utf-8", errors="replace")
            else:
                st.info("No JD files found in data/jd/")

        top_k = st.slider("Candidates to return", 1, 10, 5)

        if st.button("Run Matching", use_container_width=True,
                     disabled=not jd_text.strip()):
            with st.spinner("Running pipeline..."):
                import time
                from graphs.query_graph import query_graph
                t0 = time.time()
                try:
                    result = query_graph.invoke({
                        "jd_raw_text":  jd_text,
                        "jd_file_name": jd_fname,
                        "node_logs":    [],
                        "failed_docs":  [],
                    })
                    st.session_state["qr"]      = result
                    st.session_state["elapsed"] = time.time() - t0
                    st.session_state["top_k"]   = top_k
                except Exception as e:
                    st.error("Pipeline error: " + str(e))

    with right:
        st.markdown("<div class='shdr'>Ranked Candidates</div>", unsafe_allow_html=True)

        if "qr" in st.session_state:
            result  = st.session_state["qr"]
            final   = result.get("final_scores") or []
            parsed  = result.get("parsed_jd") or {}
            summary = result.get("results_summary") or {}
            elapsed = st.session_state.get("elapsed") or 0
            top_k   = st.session_state.get("top_k") or 5

            # JD summary panel
            if parsed:
                req_skills = [s.get("skill") or "" for s in (parsed.get("required_skills") or [])]
                chips      = "".join("<span class='chip m'>" + s + "</span>" for s in req_skills)
                title      = str(parsed.get("title") or "Untitled")
                exp_min    = str(parsed.get("experience_min_years") or 0)
                domain     = str(parsed.get("domain") or "General")
                st.markdown(
                    "<div class='jpanel'>"
                    "<div class='jtitle'>" + title + "</div>"
                    "<div class='jmeta'>" + exp_min + "y min &middot; " + domain
                    + " &middot; " + str(len(req_skills)) + " required skills</div>"
                    "<div style='margin-top:8px;'>" + chips + "</div>"
                    "</div>",
                    unsafe_allow_html=True,
                )

            t_str = str(round(elapsed, 1))
            st.markdown(
                "<div class='jmeta' style='margin-bottom:10px;'>"
                "Completed in " + t_str + "s &middot; "
                + str(len(final)) + " candidates ranked</div>",
                unsafe_allow_html=True,
            )

            if not final:
                st.warning("No candidates found. Ingest resumes first.")
            else:
                req_set = {(s.get("skill") or "").lower()
                           for s in (parsed.get("required_skills") or [])}

                for i, c in enumerate(final[:top_k]):
                    rank    = i + 1
                    name    = str(c.get("name") or "Unknown")
                    fscore  = float(c.get("final_score") or 0)
                    ats     = float(c.get("ats_score") or 0)
                    sem     = float(c.get("semantic_score") or 0)
                    rr      = float(c.get("rerank_score") or 0)
                    payload = c.get("payload") or {}
                    skills  = list(payload.get("skills") or [])
                    exp_mo  = float(payload.get("total_experience_months") or 0)
                    fname   = str(payload.get("file_name") or "")
                    edu     = payload.get("education") or []
                    degree  = edu[0].get("degree", "") if edu else ""

                    exp_yr  = str(round(exp_mo / 12, 1)) + "y"
                    fp_str  = str(round(fscore * 100)) + "%"
                    sc_col  = "#38d9a9" if fscore > 0.7 else "#5b6af5"
                    cc_cls  = "r1" if rank == 1 else ""

                    skill_chips = "".join(
                        "<span class='chip"
                        + (" m" if s.lower() in req_set else "")
                        + "'>" + s + "</span>"
                        for s in skills[:10]
                    )

                    st.markdown(
                        "<div class='ccard " + cc_cls + "'>"
                        "<div style='display:flex;justify-content:space-between;"
                        "align-items:start;'>"
                        "<div>"
                        "<div class='csub'>#" + str(rank) + " &middot; " + fname + "</div>"
                        "<div class='cname'>" + name + "</div>"
                        + ("<div class='csub' style='margin-top:2px;'>" + degree + "</div>" if degree else "")
                        + "</div>"
                        "<div style='text-align:right;'>"
                        "<div style='font-family:DM Mono,monospace;font-size:1.4rem;"
                        "font-weight:500;color:" + sc_col + ";line-height:1;'>"
                        + fp_str + "</div>"
                        "<div style='font-family:DM Mono,monospace;font-size:0.58rem;"
                        "color:#6b7280;'>FINAL SCORE</div>"
                        "</div></div>"
                        + sbar("Semantic", sem)
                        + sbar("ATS", ats, "sf2")
                        + "<div class='mrow'>"
                        + mbox(exp_yr, "Experience")
                        + mbox(str(len(skills)), "Skills")
                        + mbox(str(round(rr, 2)), "Rerank")
                        + "</div>"
                        "<div style='margin-top:8px;'>" + skill_chips + "</div>"
                        "</div>",
                        unsafe_allow_html=True,
                    )

            with st.expander("Pipeline execution log"):
                for l in (result.get("node_logs") or []):
                    st.markdown(
                        lrow(l.get("node") or "", l.get("latency_ms") or 0,
                             l.get("status") or ""),
                        unsafe_allow_html=True,
                    )

        else:
            st.markdown(
                "<div style='text-align:center;padding:50px 20px;color:#6b7280;'>"
                "<div style='font-family:DM Mono,monospace;font-size:0.73rem;"
                "letter-spacing:0.08em;'>ENTER A JD AND RUN MATCHING</div></div>",
                unsafe_allow_html=True,
            )


# ==============================================================================
#  PAGE: INGEST RESUMES
# ==============================================================================
elif page == "Ingest Resumes":
    st.markdown(
        "<div class='top-bar'><h1>Ingest Resumes</h1>"
        "<span class='badge'>INGESTION PIPELINE</span></div>",
        unsafe_allow_html=True,
    )

    col1, col2 = st.columns([1, 1.1], gap="large")

    with col1:
        st.markdown("<div class='shdr'>Upload Resumes</div>", unsafe_allow_html=True)
        st.markdown(
            "<div style='font-family:DM Mono,monospace;font-size:0.63rem;color:#6b7280;"
            "margin-bottom:8px;'>Supported: PDF, DOCX, DOC, TXT</div>",
            unsafe_allow_html=True,
        )
        uploaded = st.file_uploader(
            "Drop resume files",
            type=["pdf", "docx", "doc", "txt"],
            accept_multiple_files=True,
            label_visibility="collapsed",
        )

        st.markdown("<div class='shdr'>Or Scan Folder</div>", unsafe_allow_html=True)
        from config.settings import RESUMES_DIR
        from nodes.load_documents import scan_resume_directory
        folder_files = scan_resume_directory(RESUMES_DIR)
        st.markdown(
            "<div style='font-family:DM Mono,monospace;font-size:0.68rem;"
            "color:#6b7280;margin-bottom:8px;'>"
            + str(len(folder_files)) + " supported files found in data/resumes/</div>",
            unsafe_allow_html=True,
        )
        use_folder = st.checkbox(
            "Include " + str(len(folder_files)) + " folder files",
            disabled=not folder_files,
        )

        st.markdown(
            "<div style='font-family:DM Mono,monospace;font-size:0.63rem;"
            "color:#38d9a9;background:rgba(56,217,169,0.07);"
            "border:1px solid rgba(56,217,169,0.18);border-radius:6px;"
            "padding:8px 10px;margin:10px 0;'>"
            "Duplicate detection ON &mdash; files already indexed will be skipped."
            "</div>",
            unsafe_allow_html=True,
        )

        can_ingest = bool(uploaded or (use_folder and folder_files))
        run_ingest = st.button("Run Ingestion", use_container_width=True,
                               disabled=not can_ingest)

    with col2:
        st.markdown("<div class='shdr'>Ingestion Log</div>", unsafe_allow_html=True)

        if run_ingest:
            import time
            from graphs.ingestion_graph import ingestion_graph

            # Save uploaded files to disk
            paths = list(folder_files) if use_folder else []
            for uf in (uploaded or []):
                dest = RESUMES_DIR / uf.name
                dest.write_bytes(uf.read())
                paths.append(str(dest))

            total = len(paths)
            status_text = st.empty()
            progress_bar = st.progress(0)

            # We run the full ingestion graph and use a progress callback via
            # a simple staged approach: show progress per batch.
            # LangGraph doesn't support mid-node callbacks, so we show
            # stage-based progress (each of the 5 nodes = 20% each).
            stages = [
                (0.10, "Loading documents..."),
                (0.30, "Parsing resumes with LLM..."),
                (0.50, "Validating schemas..."),
                (0.65, "Expanding skills..."),
                (0.80, "Embedding and storing..."),
                (1.00, "Done!"),
            ]

            def update_progress(pct, msg):
                progress_bar.progress(pct)
                status_text.markdown(
                    "<div style='font-family:DM Mono,monospace;font-size:0.7rem;"
                    "color:#6b7280;'>" + msg + " (" + str(total) + " files)</div>",
                    unsafe_allow_html=True,
                )

            update_progress(0.05, "Starting pipeline...")
            t0 = time.time()

            # Use a thread to run the graph while updating the progress bar
            import threading
            result_holder = {}
            error_holder  = {}

            def run_graph():
                try:
                    result_holder["result"] = ingestion_graph.invoke({
                        "resume_file_paths": paths,
                        "node_logs":         [],
                        "failed_docs":       [],
                    })
                except Exception as e:
                    error_holder["error"] = e

            thread = threading.Thread(target=run_graph, daemon=True)
            thread.start()

            # Animate progress bar while graph runs
            stage_idx = 0
            while thread.is_alive():
                elapsed_so_far = time.time() - t0
                # Rough estimate: ~8s per 10 resumes for LLM parsing
                est_total = max(30, total * 0.8)
                auto_pct  = min(0.78, elapsed_so_far / est_total)
                if stage_idx < len(stages) - 1:
                    stage_pct, stage_msg = stages[stage_idx]
                    if auto_pct >= stage_pct:
                        stage_idx += 1
                    update_progress(auto_pct, stage_msg)
                time.sleep(1.5)

            thread.join()
            progress_bar.progress(1.0)
            status_text.empty()

            if "error" in error_holder:
                st.error("Ingestion failed: " + str(error_holder["error"]))
            else:
                res     = result_holder["result"]
                elapsed = time.time() - t0
                logs    = res.get("node_logs") or []
                failed  = res.get("failed_docs") or []
                stored     = next(
                    (l.get("stored_count", 0) for l in logs
                     if l.get("node") == "embed_and_store_node"), 0,
                )
                dupes      = next(
                    (l.get("skipped_dupes", 0) for l in logs
                     if l.get("node") == "embed_and_store_node"), 0,
                )
                checkpoint = next(
                    (l.get("skipped_count", 0) for l in logs
                     if l.get("node") == "checkpoint_filter_node"), 0,
                )
                parsed_n   = next(
                    (l.get("docs_parsed", 0) for l in logs
                     if l.get("node") == "parse_resume_node"), 0,
                )

                st.success(
                    "Ingestion complete in " + str(round(elapsed, 1)) + "s: "
                    + str(stored) + " stored, "
                    + str(checkpoint) + " skipped (already indexed), "
                    + str(dupes) + " skipped (embed-level dedup), "
                    + str(parsed_n) + " parsed, "
                    + str(len(failed)) + " failed"
                )

                st.markdown("<div class='shdr' style='margin-top:14px;'>Node Timing</div>",
                            unsafe_allow_html=True)
                for l in logs:
                    st.markdown(
                        lrow(l.get("node") or "", l.get("latency_ms") or 0,
                             l.get("status") or ""),
                        unsafe_allow_html=True,
                    )

                if failed:
                    st.markdown(
                        "<div class='shdr' style='margin-top:14px;'>Failed</div>",
                        unsafe_allow_html=True,
                    )
                    for fd in failed[:20]:  # cap at 20 to avoid huge list
                        reason = (fd.get("reason") or "")[:60]
                        st.markdown(
                            "<div class='lrow lerr'>"
                            "<span class='lnode'>" + str(fd.get("file_name") or "") + "</span>"
                            "<span>" + str(fd.get("stage") or "") + "</span>"
                            "<span>" + reason + "</span>"
                            "</div>",
                            unsafe_allow_html=True,
                        )
                    if len(failed) > 20:
                        st.markdown(
                            "<div style='font-family:DM Mono,monospace;font-size:0.63rem;"
                            "color:#6b7280;margin-top:4px;'>... and "
                            + str(len(failed) - 20) + " more</div>",
                            unsafe_allow_html=True,
                        )

                st.session_state["ir"] = res

        elif "ir" not in st.session_state:
            st.markdown(
                "<div style='text-align:center;padding:50px 20px;color:#6b7280;'>"
                "<div style='font-family:DM Mono,monospace;font-size:0.68rem;'>"
                "UPLOAD FILES OR SCAN FOLDER TO BEGIN</div></div>",
                unsafe_allow_html=True,
            )


# ==============================================================================
#  PAGE: ANALYTICS
# ==============================================================================
elif page == "Analytics":
    st.markdown(
        "<div class='top-bar'><h1>Analytics</h1>"
        "<span class='badge'>SYSTEM STATS</span></div>",
        unsafe_allow_html=True,
    )

    try:
        from config.settings import (
            get_qdrant_client, reset_qdrant_client,
            QDRANT_RESUME_COLLECTION, QDRANT_JD_COLLECTION,
            QDRANT_VECTOR_SIZE, AZURE_CHAT_DEPLOYMENT, AZURE_EMBEDDING_DEPLOYMENT,
            CROSS_ENCODER_MODEL, TOP_K_VECTOR, TOP_K_BM25, TOP_K_FINAL, RRF_K,
            ATS_WEIGHTS, FINAL_SCORE_WEIGHTS, LOG_FILE, RESUMES_DIR,
        )
        from qdrant_client.models import Distance, VectorParams

        client = get_qdrant_client()

        def get_counts():
            try:
                rc = client.get_collection(QDRANT_RESUME_COLLECTION).points_count
            except Exception:
                rc = 0
            try:
                jc = client.get_collection(QDRANT_JD_COLLECTION).points_count
            except Exception:
                jc = 0
            return rc, jc

        def reset_collection(name):
            import time as _t
            existing = [c.name for c in client.get_collections().collections]
            if name in existing:
                client.delete_collection(name)
                # Wait until Qdrant Cloud confirms the collection is gone.
                # Without this, a new client or immediate recreate can fail
                # or target the old stale collection.
                for _ in range(10):
                    _t.sleep(0.5)
                    still_there = [c.name for c in client.get_collections().collections]
                    if name not in still_there:
                        break
            # Force a fresh singleton connection so embed_and_store_node
            # sees the clean state — this is the root cause of 0 points.
            reset_qdrant_client()
            # Get fresh client for the create call
            fresh = get_qdrant_client()
            fresh.create_collection(
                collection_name=name,
                vectors_config=VectorParams(size=QDRANT_VECTOR_SIZE, distance=Distance.COSINE),
            )
            # Verify the new empty collection is visible before returning
            for _ in range(10):
                _t.sleep(0.3)
                try:
                    info = get_qdrant_client().get_collection(name)
                    if info is not None:
                        break
                except Exception:
                    pass

        r_cnt, j_cnt = get_counts()

        # ── Stat cards ────────────────────────────────────────────────────────
        c1, c2, c3, c4 = st.columns(4)
        for col, val, lbl in [
            (c1, r_cnt,              "Resumes Indexed"),
            (c2, j_cnt,              "JDs Indexed"),
            (c3, QDRANT_VECTOR_SIZE,  "Vector Dims"),
            (c4, TOP_K_FINAL,        "Top-K Results"),
        ]:
            col.markdown(
                "<div class='scard'>"
                "<div class='sval'>" + str(val) + "</div>"
                "<div class='slbl'>" + lbl + "</div>"
                "</div>",
                unsafe_allow_html=True,
            )

        # ── Collection Management ─────────────────────────────────────────────
        st.markdown("<div class='shdr'>Collection Management</div>", unsafe_allow_html=True)

        st.markdown(
            "<div style='font-family:DM Mono,monospace;font-size:0.63rem;"
            "color:#f06595;background:rgba(240,101,149,0.07);"
            "border:1px solid rgba(240,101,149,0.2);border-radius:6px;"
            "padding:8px 12px;margin-bottom:14px;'>"
            "WARNING: Deleting a collection permanently removes all indexed vectors. "
            "You will need to re-ingest all resumes afterward."
            "</div>",
            unsafe_allow_html=True,
        )

        cm1, cm2 = st.columns(2)

        with cm1:
            st.markdown(
                "<div class='jpanel'>"
                "<div class='jtitle'>resumes_index</div>"
                "<div class='jmeta' style='margin:6px 0 10px;'>"
                + str(r_cnt) + " vectors &middot; cosine similarity</div>",
                unsafe_allow_html=True,
            )
            if st.button("Delete Resume Index", key="del_resume", use_container_width=True):
                st.session_state["confirm_del_resume"] = True

            if st.session_state.get("confirm_del_resume"):
                st.warning("Are you sure? This deletes ALL " + str(r_cnt) + " resume vectors.")
                cc1, cc2 = st.columns(2)
                with cc1:
                    if st.button("Yes, Delete", key="confirm_resume_yes", use_container_width=True):
                        with st.spinner("Deleting resumes_index..."):
                            reset_collection(QDRANT_RESUME_COLLECTION)
                        st.session_state.pop("confirm_del_resume", None)
                        st.success("resumes_index deleted and recreated (0 vectors). Re-ingest resumes to populate.")
                        st.rerun()
                with cc2:
                    if st.button("Cancel", key="confirm_resume_no", use_container_width=True):
                        st.session_state.pop("confirm_del_resume", None)
                        st.rerun()

            st.markdown("</div>", unsafe_allow_html=True)

        with cm2:
            st.markdown(
                "<div class='jpanel'>"
                "<div class='jtitle'>jd_index</div>"
                "<div class='jmeta' style='margin:6px 0 10px;'>"
                + str(j_cnt) + " vectors &middot; cosine similarity</div>",
                unsafe_allow_html=True,
            )
            if st.button("Delete JD Index", key="del_jd", use_container_width=True):
                st.session_state["confirm_del_jd"] = True

            if st.session_state.get("confirm_del_jd"):
                st.warning("Are you sure? This deletes ALL " + str(j_cnt) + " JD vectors.")
                cc1, cc2 = st.columns(2)
                with cc1:
                    if st.button("Yes, Delete", key="confirm_jd_yes", use_container_width=True):
                        with st.spinner("Deleting jd_index..."):
                            reset_collection(QDRANT_JD_COLLECTION)
                        st.session_state.pop("confirm_del_jd", None)
                        st.success("jd_index deleted and recreated (0 vectors).")
                        st.rerun()
                with cc2:
                    if st.button("Cancel", key="confirm_jd_no", use_container_width=True):
                        st.session_state.pop("confirm_del_jd", None)
                        st.rerun()

            st.markdown("</div>", unsafe_allow_html=True)

        # ── Delete Both + Re-ingest ───────────────────────────────────────────
        st.markdown(
            "<div style='background:var(--surface);border:1px solid var(--border);"
            "border-radius:10px;padding:16px 18px;margin-bottom:4px;'>"
            "<div class='jtitle'>Reset Everything &amp; Re-ingest</div>"
            "<div class='jmeta' style='margin:5px 0 12px;'>"
            "Deletes both collections, then immediately re-ingests all files "
            "from data/resumes/ through the full pipeline.</div>",
            unsafe_allow_html=True,
        )

        from nodes.load_documents import scan_resume_directory
        resume_files = scan_resume_directory(RESUMES_DIR)
        st.markdown(
            "<div style='font-family:DM Mono,monospace;font-size:0.68rem;"
            "color:#38d9a9;margin-bottom:10px;'>"
            + str(len(resume_files)) + " resume files found in data/resumes/</div>",
            unsafe_allow_html=True,
        )

        if st.button("Delete All & Re-ingest from data/resumes/",
                     key="reset_all", use_container_width=True,
                     disabled=not resume_files):
            st.session_state["confirm_reset_all"] = True

        if st.session_state.get("confirm_reset_all"):
            st.error(
                "This will DELETE both indexes and re-ingest "
                + str(len(resume_files)) + " files. Cannot be undone."
            )
            ra1, ra2 = st.columns(2)
            with ra1:
                if st.button("Yes, Reset & Re-ingest", key="ra_yes", use_container_width=True):
                    st.session_state.pop("confirm_reset_all", None)

                    # Step 1: delete both collections
                    with st.spinner("Deleting indexes..."):
                        reset_collection(QDRANT_RESUME_COLLECTION)
                        reset_collection(QDRANT_JD_COLLECTION)

                    # Step 2: run ingestion pipeline with progress bar
                    import time, threading
                    from graphs.ingestion_graph import ingestion_graph

                    total = len(resume_files)
                    prog  = st.progress(0)
                    stat  = st.empty()

                    result_holder = {}
                    error_holder  = {}

                    def _run():
                        try:
                            result_holder["r"] = ingestion_graph.invoke({
                                "resume_file_paths": resume_files,
                                "node_logs":         [],
                                "failed_docs":       [],
                            })
                        except Exception as e:
                            error_holder["e"] = e

                    t = threading.Thread(target=_run, daemon=True)
                    t.start()
                    t0 = time.time()
                    stages = [
                        (0.15, "Loading documents..."),
                        (0.35, "Parsing with LLM..."),
                        (0.55, "Validating schemas..."),
                        (0.70, "Expanding skills..."),
                        (0.85, "Embedding & storing..."),
                    ]
                    si = 0
                    while t.is_alive():
                        elapsed_s = time.time() - t0
                        est       = max(30, total * 0.8)
                        pct       = min(0.9, elapsed_s / est)
                        if si < len(stages) and pct >= stages[si][0]:
                            si += 1
                        msg = stages[min(si, len(stages)-1)][1]
                        prog.progress(pct)
                        stat.markdown(
                            "<div style='font-family:DM Mono,monospace;"
                            "font-size:0.68rem;color:#6b7280;'>"
                            + msg + " (" + str(total) + " files)</div>",
                            unsafe_allow_html=True,
                        )
                        time.sleep(1.5)
                    t.join()
                    prog.progress(1.0)
                    stat.empty()

                    if "e" in error_holder:
                        st.error("Re-ingestion failed: " + str(error_holder["e"]))
                    else:
                        res    = result_holder["r"]
                        logs   = res.get("node_logs") or []
                        stored = next(
                            (l.get("stored_count", 0) for l in logs
                             if l.get("node") == "embed_and_store_node"), 0,
                        )
                        dupes  = next(
                            (l.get("skipped_dupes", 0) for l in logs
                             if l.get("node") == "embed_and_store_node"), 0,
                        )
                        failed = res.get("failed_docs") or []
                        checkpoint_n = next(
                            (l.get("skipped_count", 0) for l in logs
                             if l.get("node") == "checkpoint_filter_node"), 0,
                        )
                        st.success(
                            "Re-ingestion complete: "
                            + str(stored) + " stored, "
                            + str(checkpoint_n) + " skipped (already indexed), "
                            + str(dupes) + " skipped (dedup), "
                            + str(len(failed)) + " failed."
                        )
                        st.rerun()
            with ra2:
                if st.button("Cancel", key="ra_no", use_container_width=True):
                    st.session_state.pop("confirm_reset_all", None)
                    st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)

        st.markdown("<div class='shdr'>Models</div>", unsafe_allow_html=True)
        m1, m2 = st.columns(2)
        with m1:
            st.markdown(
                "<div class='jpanel'><div class='jtitle'>LLM Configuration</div>"
                "<div style='font-family:DM Mono,monospace;font-size:0.7rem;margin-top:9px;'>"
                "<div style='color:#6b7280;margin-bottom:2px;'>CHAT</div>"
                "<div style='margin-bottom:9px;'>" + str(AZURE_CHAT_DEPLOYMENT) + "</div>"
                "<div style='color:#6b7280;margin-bottom:2px;'>EMBEDDINGS</div>"
                "<div style='margin-bottom:9px;'>" + str(AZURE_EMBEDDING_DEPLOYMENT) + "</div>"
                "<div style='color:#6b7280;margin-bottom:2px;'>RERANKER</div>"
                "<div>" + str(CROSS_ENCODER_MODEL) + "</div>"
                "</div></div>",
                unsafe_allow_html=True,
            )
        with m2:
            st.markdown(
                "<div class='jpanel'><div class='jtitle'>Retrieval Settings</div>"
                "<div style='font-family:DM Mono,monospace;font-size:0.7rem;margin-top:9px;'>"
                "<div style='color:#6b7280;margin-bottom:2px;'>VECTOR TOP-K</div>"
                "<div style='margin-bottom:9px;'>" + str(TOP_K_VECTOR) + " candidates</div>"
                "<div style='color:#6b7280;margin-bottom:2px;'>BM25 TOP-K</div>"
                "<div style='margin-bottom:9px;'>" + str(TOP_K_BM25) + " candidates</div>"
                "<div style='color:#6b7280;margin-bottom:2px;'>RRF K</div>"
                "<div>" + str(RRF_K) + "</div>"
                "</div></div>",
                unsafe_allow_html=True,
            )

        st.markdown("<div class='shdr'>Score Weights</div>", unsafe_allow_html=True)
        w1, w2 = st.columns(2)
        with w1:
            st.markdown(
                "<div style='font-family:DM Mono,monospace;font-size:0.68rem;"
                "color:#6b7280;margin-bottom:5px;'>ATS WEIGHTS</div>",
                unsafe_allow_html=True,
            )
            for k, v in ATS_WEIGHTS.items():
                st.markdown(sbar(k, v), unsafe_allow_html=True)
        with w2:
            st.markdown(
                "<div style='font-family:DM Mono,monospace;font-size:0.68rem;"
                "color:#6b7280;margin-bottom:5px;'>FINAL SCORE WEIGHTS</div>",
                unsafe_allow_html=True,
            )
            for k, v in FINAL_SCORE_WEIGHTS.items():
                st.markdown(sbar(k, v, "sf2"), unsafe_allow_html=True)

        st.markdown("<div class='shdr'>Recent Pipeline Log</div>", unsafe_allow_html=True)
        if LOG_FILE.exists():
            import json
            raw = LOG_FILE.read_text(encoding="utf-8").strip().split("\n")[-20:]
            for line in reversed(raw):
                try:
                    e    = json.loads(line)
                    nd   = str(e.get("node") or e.get("message") or "")
                    ms   = e.get("latency_ms")
                    st_  = str(e.get("status") or "")
                    ts   = (e.get("asctime") or "")[:16]
                    err  = " lerr" if e.get("levelname") == "ERROR" else ""
                    ms_s = (str(ms) + "ms") if ms is not None else ""
                    st.markdown(
                        "<div class='lrow" + err + "'>"
                        "<span style='color:#6b7280;min-width:120px;'>" + ts + "</span>"
                        "<span class='lnode'>" + nd + "</span>"
                        "<span class='lms'>" + ms_s + "</span>"
                        "<span>" + st_ + "</span>"
                        "</div>",
                        unsafe_allow_html=True,
                    )
                except Exception:
                    st.markdown(
                        "<div class='lrow'>" + line[:120] + "</div>",
                        unsafe_allow_html=True,
                    )
        else:
            st.info("No pipeline logs yet.")

    except Exception as e:
        st.error("Analytics error: " + str(e))